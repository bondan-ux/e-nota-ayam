import base64
from datetime import datetime
import io
import json
import os
import openpyxl
from openpyxl.drawing.image import Image as OpenPyXLImage
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Mm, Pt, RGBColor
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
import streamlit as st

# 1. KONFIGURASI HALAMAN
st.set_page_config(
    page_title="Sistem Manajemen Ayam Segar",
    layout="wide",
    initial_sidebar_state="expanded",
)

logo_filename = None
for fname in ["ASTremove.PNG", "ASTremove.png", "AST.jpeg"]:
    if os.path.exists(fname):
        logo_filename = fname
        break


# --- HELPER BASE64 UNTUK BACKGROUND GAMBAR LOKAL ---
def get_base64_image(image_path):
    if os.path.exists(image_path):
        with open(image_path, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode()
    return None


bg_base64 = get_base64_image("back.jpg")


# --- HELPER GENERATOR WORD (.DOCX) NOTA KOTAK PRESISI ---
def generate_word_nota(tgl, bakul, group, items, total_bayar, logo_path):
    doc = Document()
    section = doc.sections[0]
    # KERTAS DIBAWAH PERSEGI (148mm x 105mm - ringkas & padat)
    section.page_width = Mm(148)
    section.page_height = Mm(105)
    section.top_margin = Mm(6)
    section.bottom_margin = Mm(6)
    section.left_margin = Mm(6)
    section.right_margin = Mm(6)

    tbl_hdr = doc.add_table(rows=1, cols=2)
    tbl_hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_hdr.autofit = False

    cell_l, cell_r = tbl_hdr.rows[0].cells
    cell_l.width = Mm(70)
    cell_r.width = Mm(66)

    p_l = cell_l.paragraphs[0]
    p_l.paragraph_format.space_before = Pt(0)
    p_l.paragraph_format.space_after = Pt(0)
    p_l.paragraph_format.line_spacing = 1.1

    if logo_path and os.path.exists(logo_path):
        run_img = p_l.add_run()
        run_img.add_picture(logo_path, width=Mm(12))
        p_l.add_run("  ")

    run_title = p_l.add_run("AYAM SEGAR TUMPANG\n")
    run_title.bold = True
    run_title.font.size = Pt(12)
    run_title.font.color.rgb = RGBColor(198, 40, 40)

    run_sub = p_l.add_run("Ds. Kambingan - Tumpang - Kab. Malang")
    run_sub.font.size = Pt(7.5)
    run_sub.font.color.rgb = RGBColor(90, 90, 90)

    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_r.paragraph_format.space_before = Pt(0)
    p_r.paragraph_format.space_after = Pt(0)

    r_info = p_r.add_run(f"Tanggal: {tgl}\nPembeli / Bakul: {bakul}\nGroup: {group}")
    r_info.font.size = Pt(8)

    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(2)
    p_spacer.paragraph_format.space_after = Pt(2)

    tbl_items = doc.add_table(rows=1, cols=4)
    tbl_items.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_items.autofit = False

    col_widths = [Mm(20), Mm(56), Mm(30), Mm(30)]
    headers = ["QTY / KG", "BARANG", "HARGA", "JUMLAH"]
    alignments = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT]

    hdr_cells = tbl_items.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "F0F0F0")
        set_cell_margins(hdr_cells[i], top=40, bottom=40, left=60, right=60)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = alignments[i]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h_text)
        r.bold = True
        r.font.size = Pt(8)

    for item in items:
        row_cells = tbl_items.add_row().cells
        kg_val = item["KG"]
        kg_str = f"{int(kg_val)}" if isinstance(kg_val, float) and kg_val.is_integer() else (f"{kg_val:.2f}" if isinstance(kg_val, float) else str(kg_val))
        h_str = f"Rp {item['Harga']:,.0f}".replace(",", ".")
        j_str = f"Rp {item['Jumlah']:,.0f}".replace(",", ".")

        vals = [kg_str, f" {item['Nama Barang']}", h_str, j_str]
        for i, val in enumerate(vals):
            row_cells[i].width = col_widths[i]
            set_cell_margins(row_cells[i], top=40, bottom=40, left=60, right=60)
            p = row_cells[i].paragraphs[0]
            p.alignment = alignments[i]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            r.font.size = Pt(8)

    tbl_ftr = doc.add_table(rows=1, cols=2)
    tbl_ftr.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_ftr.autofit = False

    f_left, f_right = tbl_ftr.rows[0].cells
    f_left.width = Mm(68)
    f_right.width = Mm(68)

    p_fl = f_left.paragraphs[0]
    p_fl.paragraph_format.space_before = Pt(6)
    p_fl.paragraph_format.space_after = Pt(0)
    r_sig = p_fl.add_run("Penerima,                       Hormat Kami,\n\n\n( ............................ )   ( ............................ )")
    r_sig.font.size = Pt(7.5)

    p_fr = f_right.paragraphs[0]
    p_fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_fr.paragraph_format.space_before = Pt(8)
    p_fr.paragraph_format.space_after = Pt(0)

    r_lbl = p_fr.add_run("TOTAL : ")
    r_lbl.bold = True
    r_lbl.font.size = Pt(10)

    tot_str = f"Rp {total_bayar:,.0f}".replace(",", ".")
    r_tot = p_fr.add_run(tot_str)
    r_tot.bold = True
    r_tot.font.size = Pt(14)
    r_tot.font.color.rgb = RGBColor(198, 40, 40)

    target_stream = io.BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()


# --- HELPER GENERATOR GAMBAR (PNG) NOTA KOTAK PRESISI ---
def generate_image_nota(tgl, bakul, group, items, total_bayar, logo_path):
    width = 600
    
    # Hitung tinggi dinamis berdasarkan jumlah item barang
    # Tujuannya supaya space kosong di bawah tabel tidak terlalu tinggi
    header_height = 80
    tbl_header_height = 24
    row_height = 24
    items_total_height = len(items) * row_height
    footer_height = 110
    padding = 30
    
    # Height disesuaikan secara otomatis, dengan batas minimal 360px
    calculated_height = header_height + tbl_header_height + items_total_height + footer_height + padding
    height = max(360, calculated_height)

    img = Image.new("RGB", (width, height), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    # Frame Pinggir
    draw.rectangle([12, 12, width - 12, height - 12], outline=(0, 0, 0), width=2)

    bold_fonts = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf",
        "arialbd.ttf"
    ]
    reg_fonts = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
        "arial.ttf"
    ]

    def load_font(font_list, size):
        for fn in font_list:
            try:
                return ImageFont.truetype(fn, size)
            except IOError:
                continue
        try:
            return ImageFont.load_default(size=size)
        except TypeError:
            return ImageFont.load_default()

    font_title = load_font(bold_fonts, 18)
    font_total = load_font(bold_fonts, 24)
    font_bold = load_font(bold_fonts, 13)
    font_regular = load_font(reg_fonts, 12)
    font_small = load_font(reg_fonts, 10)

    if logo_path and os.path.exists(logo_path):
        try:
            logo_img = Image.open(logo_path).convert("RGBA")
            logo_img.thumbnail((50, 50))
            img.paste(logo_img, (22, 20), logo_img)
        except Exception:
            pass

    draw.text((80, 20), "AYAM SEGAR TUMPANG", fill=(198, 40, 40), font=font_title)
    draw.text((80, 48), "Ds. Kambingan - Tumpang - Kab. Malang", fill=(90, 90, 90), font=font_small)

    draw.text((width - 200, 20), f"Tanggal: {tgl}", fill=(0, 0, 0), font=font_small)
    draw.text((width - 200, 36), f"Pembeli / Bakul: {bakul}", fill=(0, 0, 0), font=font_small)
    draw.text((width - 200, 52), f"Group: {group}", fill=(0, 0, 0), font=font_small)

    y_tbl = 80

    draw.rectangle([22, y_tbl, width - 22, y_tbl + tbl_header_height], fill=(240, 240, 240), outline=(0, 0, 0))
    draw.line([(100, y_tbl), (100, y_tbl + tbl_header_height)], fill=(0, 0, 0))
    draw.line([(330, y_tbl), (330, y_tbl + tbl_header_height)], fill=(0, 0, 0))
    draw.line([(440, y_tbl), (440, y_tbl + tbl_header_height)], fill=(0, 0, 0))

    draw.text((32, y_tbl + 4), "QTY / KG", fill=(0, 0, 0), font=font_bold)
    draw.text((110, y_tbl + 4), "BARANG", fill=(0, 0, 0), font=font_bold)
    draw.text((345, y_tbl + 4), "HARGA", fill=(0, 0, 0), font=font_bold)
    draw.text((455, y_tbl + 4), "JUMLAH", fill=(0, 0, 0), font=font_bold)

    y_curr = y_tbl + tbl_header_height
    for item in items:
        draw.rectangle([22, y_curr, width - 22, y_curr + row_height], outline=(0, 0, 0))
        draw.line([(100, y_curr), (100, y_curr + row_height)], fill=(0, 0, 0))
        draw.line([(330, y_curr), (330, y_curr + row_height)], fill=(0, 0, 0))
        draw.line([(440, y_curr), (440, y_curr + row_height)], fill=(0, 0, 0))

        kg_val = item["KG"]
        kg_str = f"{int(kg_val)}" if isinstance(kg_val, float) and kg_val.is_integer() else (f"{kg_val:.2f}" if isinstance(kg_val, float) else str(kg_val))
        h_str = f"Rp {item['Harga']:,.0f}".replace(",", ".")
        j_str = f"Rp {item['Jumlah']:,.0f}".replace(",", ".")

        draw.text((35, y_curr + 3), kg_str, fill=(0, 0, 0), font=font_regular)
        draw.text((110, y_curr + 3), item["Nama Barang"], fill=(0, 0, 0), font=font_regular)
        draw.text((340, y_curr + 3), h_str, fill=(0, 0, 0), font=font_regular)
        draw.text((450, y_curr + 3), j_str, fill=(0, 0, 0), font=font_regular)
        y_curr += row_height

    # Posisi footer sekarang dinamis mengikut baris terakhir tabel
    y_ftr = y_curr + 15

    draw.text((30, y_ftr), "Penerima,", fill=(0, 0, 0), font=font_regular)
    draw.text((160, y_ftr), "Hormat Kami,", fill=(0, 0, 0), font=font_regular)
    draw.text((25, y_ftr + 45), "( ............................ )", fill=(0, 0, 0), font=font_regular)
    draw.text((150, y_ftr + 45), "( ............................ )", fill=(0, 0, 0), font=font_regular)

    tot_str = f"Rp {total_bayar:,.0f}".replace(",", ".")
    draw.text((310, y_ftr + 15), "TOTAL :", fill=(0, 0, 0), font=font_bold)
    draw.text((380, y_ftr + 8), tot_str, fill=(198, 40, 40), font=font_total)

    img_byte_arr = io.BytesIO()
    img.save(img_byte_arr, format="PNG")
    return img_byte_arr.getvalue()


# --- CSS STYLING UTAMA ---
bg_css = ""
if bg_base64:
    bg_css = f"""
        .stApp {{
            background-image: url("data:image/jpeg;base64,{bg_base64}");
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
            background-attachment: fixed;
        }}
    """

st.markdown(
    f"""
    <style>
        footer, #MainMenu {{ visibility: hidden; }}
        [data-testid="stHeader"] {{ background-color: transparent !important; z-index: 100 !important; }}
        [data-testid="stSidebar"] {{ background-color: #FFEBEE !important; border-right: 3px solid #C62828 !important; }}
        [data-testid="stSidebar"] .stRadio label {{ font-size: 15px !important; font-weight: bold !important; color: #262626 !important; }}
        
        {bg_css}

        .block-container {{
            padding-top: 1.5rem !important;
        }}

        div[data-testid="stColumn"] > div:has(input[type="password"]) {{
            background-color: rgba(255, 255, 255, 0.98);
            padding: 30px;
            border-radius: 16px;
            box-shadow: 0 10px 30px rgba(0, 0, 0, 0.25);
            border: 1px solid #FFCDD2;
        }}
    </style>
""",
    unsafe_allow_html=True,
)

# --- 2. LOGIN ---
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.role = ""


def login():
    st.markdown("<div style='margin-top: 40px;'></div>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 1.8, 1])

    with col2:
        if logo_filename:
            c_l, c_img, c_r = st.columns([1, 2, 1])
            with c_img:
                st.image(logo_filename, use_container_width=True)

        st.markdown(
            "<h3 style='text-align: center; color: #C62828; margin-top:"
            " 10px;'>Ayam Segar Tumpang</h3>",
            unsafe_allow_html=True,
        )
        st.markdown(
            "<p style='text-align: center; color: #666; font-size: 13px;'>Silakan"
            " login untuk mengakses sistem</p>",
            unsafe_allow_html=True,
        )

        username = st.text_input("Username")
        password = st.text_input("Password", type="password")

        st.markdown("<div style='margin-top: 15px;'></div>", unsafe_allow_html=True)
        if st.button("Masuk / Login", use_container_width=True, type="primary"):
            if username == "admin" and password == "admin123":
                st.session_state.logged_in = True
                st.session_state.role = "Admin"
                st.rerun()
            elif username == "kasir" and password == "kasir123":
                st.session_state.logged_in = True
                st.session_state.role = "Kasir"
                st.rerun()
            else:
                st.error("Username atau Password salah!")


if not st.session_state.logged_in:
    login()
    st.stop()

# --- 3. MASTER HARGA JSON ---
FILE_HARGA = "master_harga.json"
default_harga = {
    "glondong": 28500,
    "jeroan": 12000,
    "usus": 16500,
    "telur_a": 269000,
    "telur_b": 250000,
    "peti": 2000,
    "box": 28500,
}

if os.path.exists(FILE_HARGA):
    try:
        with open(FILE_HARGA, "r") as f:
            saved_harga = json.load(f)
    except Exception:
        saved_harga = default_harga
else:
    saved_harga = default_harga

# --- 4. SIDEBAR ---
with st.sidebar:
    if logo_filename:
        st.image(logo_filename, width=90)
    st.markdown(
        "<h2 style='color: #C62828; margin: 0; font-size: 26px; font-weight:"
        " bold;'>AST SYSTEM</h2>",
        unsafe_allow_html=True,
    )
    st.write(f"Logged in as: **{st.session_state.role}**")
    st.markdown("---")

    st.markdown("### 📌 NAVIGASI UTAMA")
    menu_options = [
        "📊 Dashboard",
        "🧾 Nota",
        "🛍️ Penjualan",
        "📦 Stock",
        "💵 Finance",
        "⏱️ Absensi & Jadwal",
    ]
    selected_menu = st.radio("Pilih Halaman:", menu_options)

    sub_menu = None
    if selected_menu == "🧾 Nota":
        with st.expander("📂 Sub-Menu Nota", expanded=True):
            sub_menu = st.radio("Tipe Nota:", ["📑 Bakul", "🏬 Bedak", "🤝 Mitra"])

    if selected_menu == "🧾 Nota":
        st.markdown("---")
        st.markdown(
            "<h3 style='color: #C62828;'>⚙️ Master Harga</h3>",
            unsafe_allow_html=True,
        )

        def input_harga(label, key_name, default_val, step_val):
            val = st.number_input(
                f"{label}",
                value=int(saved_harga.get(key_name, default_val)),
                step=step_val,
                format="%d",
            )
            st.caption(f"➔ **Rp {val:,.0f}**".replace(",", "."))
            return val

        h_glondong = input_harga("Harga Glondong", "glondong", 28500, 500)
        h_jeroan = input_harga("Harga Jeroan", "jeroan", 12000, 500)
        h_usus = input_harga("Harga Usus", "usus", 16500, 500)
        h_telur_a = input_harga("Harga Telur A", "telur_a", 269000, 1000)
        h_telur_b = input_harga("Harga Telur B", "telur_b", 250000, 1000)
        h_peti = input_harga("Harga Peti", "peti", 2000, 100)
        h_box = input_harga("Harga Box", "box", 28500, 500)

        current_harga = {
            "glondong": h_glondong,
            "jeroan": h_jeroan,
            "usus": h_usus,
            "telur_a": h_telur_a,
            "telur_b": h_telur_b,
            "peti": h_peti,
            "box": h_box,
        }
        if current_harga != saved_harga:
            with open(FILE_HARGA, "w") as f:
                json.dump(current_harga, f)

    st.markdown("---")
    if st.button("🚪 Logout", use_container_width=True):
        st.session_state.logged_in = False
        st.session_state.role = ""
        st.rerun()

# --- 5. AREA KONTEN UTAMA ---
col_h1, col_h2 = st.columns([3, 1])
with col_h1:
    st.markdown(
        f"<h2 style='color: #C62828; margin:0;'>Ayam Segar Tumpang -"
        f" {selected_menu}</h2>",
        unsafe_allow_html=True,
    )
with col_h2:
    st.markdown(
        "<div style='text-align: right; font-weight: bold; padding-top:"
        f" 10px;'>👤 {st.session_state.role}</div>",
        unsafe_allow_html=True,
    )

st.markdown(
    "<hr style='border: 1px solid #C62828; margin-top: 5px; margin-bottom:"
    " 20px;'>",
    unsafe_allow_html=True,
)

if selected_menu == "🧾 Nota":
    if sub_menu == "📑 Bakul" or sub_menu is None:
        col_up1, col_up2 = st.columns([2, 1])
        with col_up1:
            uploaded_file = st.file_uploader(
                "Upload File Rekap Excel (.xlsx)", type=["xlsx"]
            )
        with col_up2:
            tanggal_transaksi = st.date_input(
                "Tanggal Transaksi", value=datetime.today()
            )

        if uploaded_file is not None:
            all_sheets = pd.read_excel(uploaded_file, sheet_name=None, header=None)
            valid_sheets = [
                s
                for s in all_sheets.keys()
                if s not in ["TOTAL TONASE", "NOTA FR", "Sheet1", "ploting"]
            ]

            tab_satuan, tab_bulk = st.tabs(
                ["📄 Nota Satuan (Word/PNG)", "📦 Export All Nota (Excel Filter)"]
            )

            # ==========================================
            # TAB 1: NOTA SATUAN (PILIH 1 BAKUL)
            # ==========================================
            with tab_satuan:
                selected_sheet = st.selectbox("Pilih Group / Sheet", valid_sheets)
                df_raw = all_sheets[selected_sheet]

                header_idx = 0
                for idx, row in df_raw.iterrows():
                    if row.astype(str).str.upper().str.contains("NAMA").any():
                        header_idx = idx
                        break

                df = df_raw.iloc[header_idx + 1 :].copy()
                df.columns = [
                    str(c).strip().upper() for c in df_raw.iloc[header_idx].values
                ]

                name_col = next((c for c in df.columns if "NAMA" in c), df.columns[1])
                no_col = next(
                    (c for c in df.columns if c in ["NO", "NO.", "NOMOR"]), None
                )

                df = df[df[name_col].notna()].copy()

                bakul_options = []
                bakul_map = {}

                for idx, (real_idx, row) in enumerate(df.iterrows(), start=1):
                    nama_bakul = (
                        str(row[name_col]).strip() if pd.notna(row[name_col]) else ""
                    )
                    if not nama_bakul:
                        continue

                    if (
                        no_col
                        and pd.notna(row[no_col])
                        and str(row[no_col]).strip() != ""
                    ):
                        no_val = str(row[no_col]).strip()
                        if no_val.endswith(".0"):
                            no_val = no_val[:-2]
                        label = f"{no_val}. {nama_bakul}"
                    else:
                        label = f"{idx}. {nama_bakul}"

                    bakul_options.append(label)
                    bakul_map[label] = (nama_bakul, real_idx)

                selected_bakul_label = st.selectbox("Pilih Nama Bakul", bakul_options)
                selected_bakul_info = bakul_map.get(selected_bakul_label, None)

                if selected_bakul_info:
                    selected_bakul, real_idx = selected_bakul_info

                    if (
                        "last_bakul" not in st.session_state
                        or st.session_state["last_bakul"] != selected_bakul_label
                    ):
                        st.session_state["last_bakul"] = selected_bakul_label
                        st.session_state["qty_box_val"] = 0.0
                        st.session_state["qty_telur_b_val"] = 0.0

                    row_bakul = df.loc[real_idx]

                    def get_valid_float(col_name):
                        try:
                            if not col_name or col_name not in df.columns:
                                return 0.0
                            val = row_bakul[col_name]
                            num = pd.to_numeric(val, errors="coerce")
                            return 0.0 if pd.isna(num) else float(num)
                        except Exception:
                            return 0.0

                    col_in1, col_in2, col_in3 = st.columns(3)
                    with col_in1:
                        qty_peti = st.number_input("Jumlah Peti", value=0, step=1)
                    with col_in2:
                        qty_box = st.number_input(
                            "Jumlah Box (Manual)", key="qty_box_val", step=1
                        )
                    with col_in3:
                        qty_telur_b_manual = st.number_input(
                            "Jumlah Telur B (Manual)", key="qty_telur_b_val", step=1
                        )

                    qty_tonase = get_valid_float(
                        next((c for c in df.columns if "TONASE" in c), "")
                    )
                    qty_jeroan = get_valid_float(
                        next((c for c in df.columns if "JEROAN" in c), "")
                    )
                    qty_usus = get_valid_float(
                        next((c for c in df.columns if "USUS" in c), "")
                    )
                    qty_telur_a = get_valid_float(
                        next(
                            (c for c in df.columns if "TELUR A" in c or "TELUR" in c), ""
                        )
                    )

                    qty_telur_b_excel = get_valid_float(
                        next((c for c in df.columns if "TELUR B" in c), "")
                    )
                    qty_telur_b = (
                        qty_telur_b_excel
                        if qty_telur_b_excel > 0
                        else qty_telur_b_manual
                    )

                    val_ket = get_valid_float(
                        next((c for c in df.columns if "KET" in c), "")
                    )
                    biaya_kresek = (
                        7000 if (val_ket > 0 and not float(val_ket).is_integer()) else 0
                    )

                    tot_glondong = qty_tonase * h_glondong
                    tot_jeroan = qty_jeroan * h_jeroan
                    tot_usus = qty_usus * h_usus
                    tot_telur_a = qty_telur_a * h_telur_a
                    tot_telur_b = qty_telur_b * h_telur_b
                    tot_peti = qty_peti * h_peti
                    tot_box = qty_box * h_box

                    total_bayar = (
                        tot_glondong
                        + tot_jeroan
                        + tot_usus
                        + tot_telur_a
                        + tot_telur_b
                        + tot_peti
                        + tot_box
                        + biaya_kresek
                    )

                    items = [
                        {
                            "Nama Barang": "GLONDONG",
                            "KG": qty_tonase,
                            "Harga": h_glondong,
                            "Jumlah": tot_glondong,
                        },
                        {
                            "Nama Barang": "JEROAN",
                            "KG": qty_jeroan,
                            "Harga": h_jeroan,
                            "Jumlah": tot_jeroan,
                        },
                        {
                            "Nama Barang": "USUS B",
                            "KG": qty_usus,
                            "Harga": h_usus,
                            "Jumlah": tot_usus,
                        },
                        {
                            "Nama Barang": "TELUR A",
                            "KG": qty_telur_a,
                            "Harga": h_telur_a,
                            "Jumlah": tot_telur_a,
                        },
                        {
                            "Nama Barang": "TELUR B",
                            "KG": qty_telur_b,
                            "Harga": h_telur_b,
                            "Jumlah": tot_telur_b,
                        },
                        {
                            "Nama Barang": "PETI",
                            "KG": qty_peti,
                            "Harga": h_peti,
                            "Jumlah": tot_peti,
                        },
                        {
                            "Nama Barang": "BOX",
                            "KG": qty_box,
                            "Harga": h_box,
                            "Jumlah": tot_box,
                        },
                        {
                            "Nama Barang": "BIAYA KRESEK",
                            "KG": 1 if biaya_kresek > 0 else 0,
                            "Harga": 7000,
                            "Jumlah": biaya_kresek,
                        },
                    ]

                    filtered_items = [i for i in items if i["KG"] > 0]

                    if filtered_items:
                        with st.container(border=True):
                            col_l, col_m, col_r = st.columns([1.5, 4, 3])
                            with col_l:
                                if logo_filename:
                                    st.image(logo_filename, width=110)
                            with col_m:
                                st.markdown(
                                    "<h3 style='margin:0; color:#C62828; font-weight: bold;"
                                    " font-size:24px;'>AYAM SEGAR TUMPANG</h3><small>Ds. Kambingan"
                                    " - Tumpang - Kab. Malang</small>",
                                    unsafe_allow_html=True,
                                )
                            with col_r:
                                st.markdown(
                                    f"<div style='text-align:right;'><small><b>Tanggal:</b>"
                                    f" {tanggal_transaksi.strftime('%d-%m-%Y')}<br><b>Bakul:</b>"
                                    f" {selected_bakul}<br><b>Group:</b>"
                                    f" {selected_sheet}</small></div>",
                                    unsafe_allow_html=True,
                                )

                            rows_html = ""
                            for item in filtered_items:
                                kg_str = (
                                    f"{int(item['KG'])}"
                                    if isinstance(item["KG"], float) and item["KG"].is_integer()
                                    else (
                                        f"{item['KG']:.2f}"
                                        if isinstance(item["KG"], float)
                                        else str(item["KG"])
                                    )
                                )
                                rows_html += f"<tr><td style='text-align: center;'>{kg_str}</td><td>{item['Nama Barang']}</td><td style='text-align: right;'>Rp {item['Harga']:,.0f}</td><td style='text-align: right;'>Rp {item['Jumlah']:,.0f}</td></tr>".replace(
                                    ",", "."
                                )

                            st.markdown(
                                f"""
                                <table style="width:100%; border-collapse:collapse; margin-top:10px;">
                                    <thead><tr><th style="border:1px solid #000; padding:6px; background:#f2f2f2;">QTY / KG</th><th style="border:1px solid #000; padding:6px; background:#f2f2f2;">BARANG</th><th style="border:1px solid #000; padding:6px; background:#f2f2f2;">HARGA</th><th style="border:1px solid #000; padding:6px; background:#f2f2f2;">JUMLAH</th></tr></thead>
                                    <tbody>{rows_html}</tbody>
                                </table>
                                <div style="text-align:right; margin-top:10px; font-weight:bold;">
                                    TOTAL: <span style="color:#C62828; font-size:22px; font-weight:bold;">Rp {total_bayar:,.0f}</span>
                                </div>
                            """.replace(",", "."),
                                unsafe_allow_html=True,
                            )

                        col_btn1, col_btn2 = st.columns(2)

                        with col_btn1:
                            word_data = generate_word_nota(
                                tgl=tanggal_transaksi.strftime("%d-%m-%Y"),
                                bakul=selected_bakul,
                                group=selected_sheet,
                                items=filtered_items,
                                total_bayar=total_bayar,
                                logo_path=logo_filename,
                            )
                            st.download_button(
                                label="📝 Download Nota Word (.docx)",
                                data=word_data,
                                file_name=(
                                    f"Nota_{selected_bakul}_{tanggal_transaksi.strftime('%Y%m%d')}.docx"
                                ),
                                mime=(
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                ),
                                use_container_width=True,
                            )

                        with col_btn2:
                            img_data = generate_image_nota(
                                tgl=tanggal_transaksi.strftime("%d-%m-%Y"),
                                bakul=selected_bakul,
                                group=selected_sheet,
                                items=filtered_items,
                                total_bayar=total_bayar,
                                logo_path=logo_filename,
                            )
                            st.download_button(
                                label="🖼️ Download Nota Gambar (.png)",
                                data=img_data,
                                file_name=(
                                    f"Nota_{selected_bakul}_{tanggal_transaksi.strftime('%Y%m%d')}.png"
                                ),
                                mime="image/png",
                                use_container_width=True,
                            )

            # ==========================================
            # TAB 2: EXPORT ALL NOTA (EMBED GAMBAR KE EXCEL)
            # ==========================================
            with tab_bulk:
                st.markdown("### 📄 Export All Nota ke File Excel (Format Gambar Nota)")
                st.caption(
                    "Centang nama bakul yang ingin diproses. Gambar nota JPG akan disusun"
                    " rapi di dalam sheet Excel."
                )

                selected_bakul_by_sheet = {}

                for sheet_name in valid_sheets:
                    df_raw_sheet = all_sheets[sheet_name]

                    header_idx_sheet = 0
                    for idx, row in df_raw_sheet.iterrows():
                        if row.astype(str).str.upper().str.contains("NAMA").any():
                            header_idx_sheet = idx
                            break

                    df_s = df_raw_sheet.iloc[header_idx_sheet + 1 :].copy()
                    df_s.columns = [
                        str(c).strip().upper()
                        for c in df_raw_sheet.iloc[header_idx_sheet].values
                    ]

                    name_c = next(
                        (c for c in df_s.columns if "NAMA" in c), df_s.columns[1]
                    )
                    df_s = df_s[df_s[name_c].notna()].copy()

                    list_bakul = [
                        str(b).strip() for b in df_s[name_c].tolist() if str(b).strip()
                    ]

                    if list_bakul:
                        with st.expander(
                            f"📋 Sheet: {sheet_name} ({len(list_bakul)} Bakul)",
                            expanded=True,
                        ):
                            col_a, col_b = st.columns([1, 4])
                            with col_a:
                                select_all = st.checkbox(
                                    "Pilih Semua", value=True, key=f"all_xl_img_{sheet_name}"
                                )

                            default_selected = list_bakul if select_all else []
                            chosen = st.multiselect(
                                "Daftar Bakul Terpilih:",
                                options=list_bakul,
                                default=default_selected,
                                key=f"ms_xl_img_{sheet_name}",
                            )
                            selected_bakul_by_sheet[sheet_name] = (df_s, name_c, chosen)

                st.markdown("---")

                if st.button(
                    "🚀 Generate & Download Excel All Nota Gambar",
                    type="primary",
                    use_container_width=True,
                ):
                    wb = openpyxl.Workbook()
                    wb.remove(wb.active)

                    total_processed = 0

                    for sheet_name, (
                        df_s,
                        name_c,
                        chosen_bakul,
                    ) in selected_bakul_by_sheet.items():
                        if not chosen_bakul:
                            continue

                        ws = wb.create_sheet(title=sheet_name[:31])
                        df_filtered = df_s[
                            df_s[name_c].astype(str).str.strip().isin(chosen_bakul)
                        ].copy()

                        row_position = 2

                        for idx, row in df_filtered.iterrows():
                            nama_bakul = str(row[name_c]).strip()

                            def get_val(c_name):
                                try:
                                    col = next((c for c in df_filtered.columns if c_name in c), "")
                                    val = pd.to_numeric(row[col], errors="coerce")
                                    return 0.0 if pd.isna(val) else float(val)
                                except Exception:
                                    return 0.0

                            kg_tonase = get_val("TONASE")
                            kg_jeroan = get_val("JEROAN")
                            kg_usus = get_val("USUS")
                            kg_telur_a = get_val("TELUR A")
                            kg_telur_b = get_val("TELUR B")
                            val_ket = get_val("KET")
                            biaya_kresek = (
                                7000 if (val_ket > 0 and not float(val_ket).is_integer()) else 0
                            )

                            tot_glondong = kg_tonase * h_glondong
                            tot_jeroan = kg_jeroan * h_jeroan
                            tot_usus = kg_usus * h_usus
                            tot_telur_a = kg_telur_a * h_telur_a
                            tot_telur_b = kg_telur_b * h_telur_b

                            tot_bayar = (
                                tot_glondong
                                + tot_jeroan
                                + tot_usus
                                + tot_telur_a
                                + tot_telur_b
                                + biaya_kresek
                            )

                            items = [
                                {
                                    "Nama Barang": "GLONDONG",
                                    "KG": kg_tonase,
                                    "Harga": h_glondong,
                                    "Jumlah": tot_glondong,
                                },
                                {
                                    "Nama Barang": "JEROAN",
                                    "KG": kg_jeroan,
                                    "Harga": h_jeroan,
                                    "Jumlah": tot_jeroan,
                                },
                                {
                                    "Nama Barang": "USUS B",
                                    "KG": kg_usus,
                                    "Harga": h_usus,
                                    "Jumlah": tot_usus,
                                },
                                {
                                    "Nama Barang": "TELUR A",
                                    "KG": kg_telur_a,
                                    "Harga": h_telur_a,
                                    "Jumlah": tot_telur_a,
                                },
                                {
                                    "Nama Barang": "TELUR B",
                                    "KG": kg_telur_b,
                                    "Harga": h_telur_b,
                                    "Jumlah": tot_telur_b,
                                },
                                {
                                    "Nama Barang": "BIAYA KRESEK",
                                    "KG": 1 if biaya_kresek > 0 else 0,
                                    "Harga": 7000,
                                    "Jumlah": biaya_kresek,
                                },
                            ]

                            filtered_items = [i for i in items if i["KG"] > 0]

                            if filtered_items:
                                img_bytes = generate_image_nota(
                                    tgl=tanggal_transaksi.strftime("%d-%m-%Y"),
                                    bakul=nama_bakul,
                                    group=sheet_name,
                                    items=filtered_items,
                                    total_bayar=tot_bayar,
                                    logo_path=logo_filename,
                                )

                                img_obj = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                                img_jpg_stream = io.BytesIO()
                                img_obj.save(img_jpg_stream, format="JPEG", quality=95)
                                img_jpg_stream.seek(0)

                                xl_img = OpenPyXLImage(img_jpg_stream)
                                cell_location = f"B{row_position}"
                                ws.add_image(xl_img, cell_location)

                                row_position += 28
                                total_processed += 1

                    if total_processed > 0:
                        output_excel = io.BytesIO()
                        wb.save(output_excel)

                        st.success(
                            f"✅ Berhasil menyusun {total_processed} gambar nota ke dalam"
                            " file Excel!"
                        )
                        st.download_button(
                            label="📥 Download File Excel All Nota Gambar (.xlsx)",
                            data=output_excel.getvalue(),
                            file_name=(
                                "EXCEL_NOTA_GAMBAR_"
                                f"{tanggal_transaksi.strftime('%Y%m%d')}.xlsx"
                            ),
                            mime=(
                                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                            ),
                            use_container_width=True,
                        )
                    else:
                        st.warning(
                            "⚠️ Tidak ada nota yang diproses. Pastikan bakul yang dipilih"
                            " memiliki data transaksi."
                        )

        else:
            st.info("💡 Silakan upload file Rekap Excel untuk memproses nota.")
    else:
        st.info(f"Fitur untuk Nota {sub_menu} siap dikembangkan.")
else:
    st.info(f"Halaman {selected_menu} sedang dalam pengembangan.")
